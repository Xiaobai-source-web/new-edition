"""调用 Dify Chatflow 工作流 API。

增强点：
1. API 调用重试机制（偶发网络波动自动重试，业务错误不重试）
2. 节点进度回调（实时通知调用方当前执行到哪个节点，用于 UI 进度反馈）
3. 完整事件处理：PING / message_replace / agent_message / node_finished 等
4. 工作流卡住（PENDING / 人工介入）时的超时提示
"""

from __future__ import annotations

import json
import re
import time
from typing import Any, Callable, Dict, List, Tuple

import requests

from config import (
    DIFY_API_KEY,
    DIFY_CHATFLOW_URL,
    DIFY_MAX_RETRIES,
    DIFY_RETRY_INTERVAL,
    DIFY_TIMEOUT,
)

# 节点标题 → 用户可读的进度描述
# 工作流执行到对应节点时，回调会传这段文字给 UI 显示
_NODE_TITLE_MAP: Dict[str, str] = {
    "1.1综合参数提取": "正在提取项目参数...",
    "1.2边界条件补充": "正在补充边界条件...",
    "1.3WBS一二级+粗颗粒逻辑": "正在生成工作分解结构（WBS）...",
    "1.4WBS三级生成": "正在细化 WBS 三级任务...",
    "1.5WBS审查": "正在审查 WBS 结构...",
    "2.1工序依赖关系": "正在计算工序依赖关系...",
    "3.1资源调优": "正在进行资源调优...",
    "3.2进度方案生成": "正在生成进度方案...",
    "4.1方案审核": "正在审核方案逻辑...",
    "4.2方案重生成": "正在重新生成方案...",
    "5.1人员配置": "正在计算人员配置...",
    "5.2资源荷载": "正在计算设备资源荷载...",
    "5.3横道图": "正在整理横道图数据...",
    "6.1扰动事件解析": "正在解析扰动事件...",
    "6.2生成更新方案": "正在生成更新方案...",
    "6.3进度方案生成": "正在生成优化后的进度方案...",
    "可视化": "正在生成方案摘要...",
    "LLM 19": "正在输出最终方案...",
    "输出整理": "正在整理输出...",
    "输出整理2": "正在整理输出...",
    "输出整理4": "正在整理输出...",
    "输出整理5": "正在整理输出...",
}


def call_dify_chatflow(
    query_text: str,
    *,
    files: list | None = None,
    conversation_id: str | None = None,
    current_plan_json: Dict[str, Any] | None = None,
    on_progress: Callable[[str], None] | None = None,
) -> Tuple[str, Dict[str, Any], str | None]:
    """调用 Dify Chatflow API（流式），返回 (文字描述, structured_output, conversation_id)。

    - 文字描述：在聊天框显示给用户（可复制）
    - structured_output：如果有，就是 {overview, all_tasks_schedule, ...} 的 JSON 字典，
      直接交给 normalize_to_wrapped + validate_data_structure 后就可以渲染图表；
      若没有可用结构则返回空字典 {}。
    - conversation_id：Dify 会话 ID，续问时要传入。
    - on_progress：可选的进度回调函数，工作流每执行到一个节点就会调用一次，
      传入该节点的用户可读描述（如 "正在生成进度方案..."）。
      UI 层可以用这个来实时显示 AI 的执行进度。
    """
    headers = {
        "Authorization": f"Bearer {DIFY_API_KEY}",
        "Content-Type": "application/json",
    }

    final_query = query_text
    if files:
        for file_obj in files:
            try:
                name = getattr(file_obj, "name", None) or "上传文件"
                if name.lower().endswith(".docx"):
                    try:
                        from docx import Document
                        from io import BytesIO

                        doc = Document(BytesIO(file_obj.read()))
                        parts: list[str] = []
                        for p in doc.paragraphs:
                            if p.text.strip():
                                parts.append(p.text)
                        for t in doc.tables:
                            for row in t.rows:
                                cells = [cell.text.strip() for cell in row.cells]
                                parts.append(" | ".join(cells))
                        content = "\n".join(parts)
                    except Exception:
                        content = f"(无法解析 .docx 文件：{name})"
                else:
                    raw = file_obj.read()
                    try:
                        content = raw.decode("utf-8") if isinstance(raw, (bytes, bytearray)) else str(raw)
                    except Exception:
                        content = f"(无法读取文件：{name})"
                final_query += f"\n\n===== 文件 {name} 内容 =====\n{content}\n"
            except Exception:
                final_query += f"\n\n(文件读取失败)\n"

    if current_plan_json:
        plan_str = json.dumps(current_plan_json, ensure_ascii=False, indent=2)
        final_query = (
            f"【用户要求】\n{final_query}\n\n"
            f"【当前已有进度计划数据】\n{plan_str}\n\n"
            f"请基于当前进度计划数据，结合用户要求进行生成/调整。"
        )

    payload: Dict[str, Any] = {
        "inputs": {},
        "query": final_query,
        "response_mode": "streaming",
        "user": "streamlit_web",
    }
    if conversation_id:
        payload["conversation_id"] = conversation_id

    # ====== 带重试的调用 ======
    last_error: Exception | None = None
    for attempt in range(1, DIFY_MAX_RETRIES + 1):
        try:
            return _do_stream_request(
                headers=headers,
                payload=payload,
                on_progress=on_progress,
            )
        except _BusinessError as exc:
            # 业务错误（AI 返回 error 事件）不重试，直接抛出
            raise RuntimeError(str(exc)) from exc
        except (requests.ConnectionError, requests.Timeout, requests.HTTPError) as exc:
            last_error = exc
            if attempt < DIFY_MAX_RETRIES:
                # 通知 UI 正在重试
                if on_progress:
                    on_progress(f"网络异常，{DIFY_RETRY_INTERVAL} 秒后第 {attempt + 1} 次尝试...")
                time.sleep(DIFY_RETRY_INTERVAL)
                continue
            # 重试用完，抛出友好错误
            raise RuntimeError(
                f"AI 服务连接失败（已重试 {DIFY_MAX_RETRIES} 次）：{exc}。请检查网络后重试。"
            ) from exc
        except Exception as exc:
            # 其他未知错误不重试
            raise RuntimeError(f"调用 AI 服务时发生错误：{exc}") from exc

    # 理论上不会走到这里
    raise RuntimeError(f"AI 服务调用失败：{last_error}")


class _BusinessError(Exception):
    """业务错误（AI 返回 error 事件），不重试。"""


def _do_stream_request(
    *,
    headers: dict,
    payload: dict,
    on_progress: Callable[[str], None] | None,
) -> Tuple[str, Dict[str, Any], str | None]:
    """执行一次流式请求，返回 (文字描述, structured_output, conversation_id)。

    遇到网络错误/超时抛出对应异常（由上层重试）；遇到业务 error 事件抛出 _BusinessError。
    当工作流触发人工介入（PENDING）导致卡住时，在 answer_text 中给出明确提示。
    """
    answer_parts: list[str] = []
    structured_output: Dict[str, Any] = {}
    new_conversation_id: str | None = None
    last_node_title: str = ""
    last_progress_time = time.time()
    got_message_end = False

    with requests.post(
        DIFY_CHATFLOW_URL,
        headers=headers,
        json=payload,
        stream=True,
        timeout=DIFY_TIMEOUT,
    ) as resp:
        resp.raise_for_status()
        for raw_line in resp.iter_lines():
            if not raw_line:
                # 空行时检查是否长时间卡住（超过 180 秒没事件就当卡住了）
                if last_node_title and time.time() - last_progress_time > 180:
                    answer_parts.append(
                        f"\n\n⚠️ 工作流执行时间过长，已停留在「{last_node_title}」阶段超过 3 分钟。"
                        " 可能原因：① Dify 工作流触发了人工介入节点，请在 Dify 后台处理后重试；"
                        " ② 项目参数过于复杂，建议拆分成多个小步骤提交。"
                    )
                    break
                continue
            try:
                line = raw_line.decode("utf-8")
            except Exception:
                continue
            if not line.startswith("data: "):
                continue
            body = line[len("data: "):]
            if body.strip() == "[DONE]":
                break
            try:
                evt = json.loads(body)
            except json.JSONDecodeError:
                continue
            event = evt.get("event")

            if event == "message":
                # 增量推送的文字片段
                answer_parts.append(evt.get("answer", ""))

            elif event == "message_replace":
                # Dify 有时会用整段新文本替换之前返回的内容
                new_text = evt.get("answer") or ""
                answer_parts = [new_text]

            elif event == "agent_message":
                # Agent 模式下的文字输出
                answer_parts.append(evt.get("answer", ""))

            elif event == "message_end":
                got_message_end = True
                new_conversation_id = evt.get("conversation_id")
                metadata = evt.get("metadata") or {}
                # Dify 工作流 answer 节点中若有结构化输出，优先从 outputs 拿
                outputs = metadata.get("outputs") if isinstance(metadata, dict) else None
                if not outputs:
                    outputs = evt.get("outputs") or {}
                if isinstance(outputs, dict):
                    # 兼容某些节点直接把 fields 放进 outputs（如 structured_output / plan）
                    for key in ("structured_output", "plan", "output"):
                        if key in outputs and isinstance(outputs[key], dict):
                            structured_output = outputs[key]
                            break
                    # 兼容：outputs 本身就是 structured_output（无外层 key）
                    if not structured_output and (
                        "overview" in outputs and "all_tasks_schedule" in outputs
                    ):
                        structured_output = outputs
                    if not structured_output and "text" in outputs:
                        candidate = _try_parse_structured_from_text(str(outputs["text"]))
                        if candidate:
                            structured_output = candidate
                if not structured_output and answer_parts:
                    candidate = _try_parse_structured_from_text("".join(answer_parts))
                    if candidate:
                        structured_output = candidate

            elif event in ("node_started", "workflow_node_started"):
                # 节点开始事件：通知 UI 当前执行到哪个节点
                if on_progress:
                    node_data = evt.get("data") or {}
                    node_title = node_data.get("title") or ""
                    if not node_title:
                        # 有些版本的 Dify 把 title 放在顶层
                        node_title = evt.get("node_title") or ""
                    last_node_title = node_title
                    last_progress_time = time.time()
                    progress_text = _NODE_TITLE_MAP.get(node_title)
                    if progress_text:
                        on_progress(progress_text)
                    else:
                        # 未知节点也显示原文，让用户知道在执行什么
                        if node_title:
                            on_progress(f"正在执行：{node_title}...")

            elif event in ("node_finished", "workflow_node_finished"):
                last_progress_time = time.time()

            elif event == "workflow_started":
                last_progress_time = time.time()
                if on_progress:
                    on_progress("AI 开始处理您的请求...")

            elif event == "workflow_finished":
                last_progress_time = time.time()
                if on_progress:
                    on_progress("处理完成，正在整理结果...")
                # 兼容某些 Dify 版本：workflow_finished 里可能带 outputs
                wf_data = evt.get("data") or {}
                if isinstance(wf_data, dict) and not structured_output:
                    outputs = wf_data.get("outputs") or {}
                    if isinstance(outputs, dict):
                        for key in ("structured_output", "plan", "output"):
                            if key in outputs and isinstance(outputs[key], dict):
                                structured_output = outputs[key]
                                break

            elif event == "ping":
                # 心跳事件，刷新最后活跃时间
                last_progress_time = time.time()

            elif event == "error":
                msg = evt.get("message", "未知错误")
                # 遇到错误时先尝试从已有的 answer_parts 中提取 structured_output
                if answer_parts and not structured_output:
                    candidate = _try_parse_structured_from_text("".join(answer_parts))
                    if candidate:
                        structured_output = candidate
                raise _BusinessError(f"Dify 返回错误：{msg}")

            # 其他未识别事件：不做处理，但刷新时间（避免被当成卡住）
            else:
                last_progress_time = time.time()

    answer_text = "".join(answer_parts).strip()

    # ====== 防御性清洗（多层、多策略）：把 answer_text 中的 JSON 彻底剔除 ======
    # 即使 Dify 代码节点把 structured_output 漏进了 final_text，这里也会剥掉，
    # 保证对话框里只给用户看中文，不显示 JSON。
    if answer_text:
        cleaned = answer_text

        # ============ 第 1 层：删除 ```json ... ``` 等各种代码块 ============
        # 1.1) 标准 ```json { ... } ``` 代码块
        cleaned = re.sub(
            r"```(?:json|JSON|structured_output)?\s*\{[\s\S]*?\}\s*```",
            "",
            cleaned,
            flags=re.IGNORECASE,
        )
        # 1.2) 任意 ```xxx ... ``` 代码块（可能包裹 JSON）
        cleaned = re.sub(r"```[\w\s-]*\s*[\s\S]*?\s*```", "", cleaned, flags=re.IGNORECASE)
        # 1.3) 不闭合的 ```json 开头到文末
        cleaned = re.sub(r"```(?:json)?\s*[\s\S]*$", "", cleaned, flags=re.IGNORECASE)

        # ============ 第 2 层：用深度匹配删除所有「看起来像进度计划」的大型 JSON 对象 ============
        #    （不仅删末尾的，全文扫描任何位置的）
        def _strip_all_plan_json(s: str) -> str:
            """全文扫描，删除所有大型进度计划 JSON 对象。"""
            result_parts = []
            i = 0
            n = len(s)
            while i < n:
                c = s[i]
                if c != "{":
                    result_parts.append(c)
                    i += 1
                    continue
                # 尝试解析以 i 开头的 JSON 对象（深度匹配）
                depth = 0
                end = -1
                in_str = False
                esc = False
                for j in range(i, n):
                    ch = s[j]
                    if esc:
                        esc = False
                        continue
                    if ch == "\\":
                        esc = True
                        continue
                    if ch == '"':
                        in_str = not in_str
                        continue
                    if in_str:
                        continue
                    if ch == "{":
                        depth += 1
                    elif ch == "}":
                        depth -= 1
                        if depth == 0:
                            end = j
                            break
                if end != -1:
                    candidate = s[i : end + 1]
                    candidate_len = len(candidate)
                    # 识别策略：大型（>200）且含关键字段的 JSON 一律视为进度计划数据
                    plan_keywords = (
                        # --- 生成场景字段 ---
                        "overview",
                        "all_tasks_schedule",
                        "structured_output",
                        "project_name",
                        "planned_start_date",
                        "planned_end_date",
                        "critical_path_tasks",
                        "key_milestones",
                        "resource_plan",
                        "section_code",
                        "assigned_resources",
                        # --- 优化/调整场景外层 key ---
                        "adjusted_plan",
                        "optimized_plan",
                        "updated_plan",
                        "updated_schedule",
                        "adjusted_schedule",
                        "optimized_schedule",
                        "new_schedule",
                        "final_plan",
                        "delay_adjusted_plan",
                        # --- 优化场景中文描述（常作为key）---
                        "优化后计划",
                        "调整后进度",
                        "延误调整计划",
                        "压缩后计划",
                    )
                    hit_count = sum(1 for kw in plan_keywords if kw in candidate)
                    if candidate_len > 200 and hit_count >= 2:
                        # 跳过这个 JSON，不写入 result_parts
                        i = end + 1
                        # JSON 前后的多余空白/换行/逗号一并吞掉
                        while i < n and s[i] in " \t\r\n,，;；":
                            i += 1
                        # 回退：result_parts 末尾的多余空白也去掉
                        while result_parts and result_parts[-1] in " \t\r\n":
                            result_parts.pop()
                        if result_parts:
                            result_parts.append("\n")
                        continue
                    else:
                        # 小 JSON 或不是进度结构的，保留原字符
                        result_parts.append(c)
                        i += 1
                else:
                    result_parts.append(c)
                    i += 1
            return "".join(result_parts)

        cleaned = _strip_all_plan_json(cleaned)

        # ============ 第 3 层：兜底清理「{structured_output: ...}」这种外层包裹 ============
        #    以及一些残留的 JSON 碎片（如 "structured_output": {...} 片段）
        def _strip_structured_wrapper(s: str) -> str:
            """删除 structured_output 字段形式的 JSON 碎片。"""
            pattern = r'"structured_output"\s*:\s*\{[\s\S]*\}\s*'
            return re.sub(pattern, "", s, flags=re.DOTALL)

        cleaned = _strip_structured_wrapper(cleaned)

        # ============ 第 4 层：清理残留的 JSON 标签行（如 "以下是JSON数据：" 等引导语） ============
        lines = cleaned.splitlines()
        filtered_lines = []
        skip_keywords = (
            # --- JSON / 结构化输出相关 ---
            "json",
            "JSON",
            "进度计划数据",
            "结构化输出",
            "structured_output",
            "all_tasks_schedule",
            "adjusted_plan",
            "optimized_plan",
            "updated_plan",
            # --- 通用引导语 ---
            "下面是",
            "如下为",
            "以下为",
            "请复制",
            "请保存为",
            "```",
            # --- 优化场景引导语 ---
            "优化后计划如下",
            "调整后计划如下",
            "延误调整后计划",
            "压缩后计划",
            "更新的进度计划",
            "方案数据如下",
            "绘图数据",
            "供系统绘图使用",
            "系统内部使用",
            "数据结构",
            "以下JSON",
            "详细数据",
            "数据内容",
            "JSON格式",
            "json格式",
            "仅供系统",
            "供绘图使用",
        )
        for line in lines:
            stripped = line.strip()
            if not stripped:
                # 连续空行最多保留 1 个
                if filtered_lines and filtered_lines[-1] == "":
                    continue
                filtered_lines.append("")
                continue
            # 如果一行包含大量 { } [ ] " : , 这类 JSON 字符，直接跳过
            json_chars = sum(stripped.count(ch) for ch in '{}[]":,')
            if json_chars > len(stripped) * 0.2:  # JSON 字符占比超过 20%
                continue
            # 行内有关键词且行很短（引导语），跳过
            lower = stripped.lower()
            hit = any(kw.lower() in lower for kw in skip_keywords)
            if hit and len(stripped) < 60:
                continue
            filtered_lines.append(line)
        cleaned = "\n".join(filtered_lines)

        # ============ 最终 trim ============
        cleaned = cleaned.strip()
        if cleaned:
            answer_text = cleaned

    # 如果没收到 message_end，说明流式被异常中断或人工介入卡住了
    if not got_message_end and not answer_text:
        stuck_msg = (
            "⚠️ AI 工作流未正常结束。可能原因：\n"
            "① Dify 工作流触发了人工介入节点（PENDING 状态），请在 Dify 后台查看日志并处理；\n"
            "② 工作流某个节点执行出错，检查代码节点或 LLM Prompt；\n"
            "③ 工作流最新版本未发布，API 仍在调用旧版本，请确认 Dify 编辑器右上角已点「发布」。"
        )
        answer_text = stuck_msg

    return answer_text, structured_output, new_conversation_id


def _try_parse_structured_from_text(text: str) -> Dict[str, Any]:
    """尝试从 AI 回答文本里抓取符合绘图结构的 JSON。

    搜索顺序：
    1) ```json ... ``` 代码块
    2) 全文扫描，对每个完整 JSON 对象尝试解析 + 解包外层包装

    支持多种外层包装：{structured_output, adjusted_plan, optimized_plan, ...}
    最终返回统一格式：{overview, all_tasks_schedule, ...}（扁平结构），
    上层 normalize_to_wrapped 会再把它包装成 {structured_output: ...}。
    """
    if not text:
        return {}

    import re

    def _unwrap_any_plan(data: Any) -> Dict[str, Any]:
        """从任意形式（含多种外层 key / 多层嵌套）的 data 中解出实际的 plan dict。

        返回的 dict 直接包含 overview / all_tasks_schedule（扁平形式）。
        如果解不出来，返回空 dict {}。
        """
        if not isinstance(data, dict):
            return {}

        # 形式 A：data 本身就是扁平计划
        if "overview" in data and "all_tasks_schedule" in data:
            return data

        # 形式 B：含多种常见外层 key
        outer_keys = (
            "structured_output",
            "plan",
            "output",
            "result",
            "adjusted_plan",
            "optimized_plan",
            "updated_plan",
            "updated_schedule",
            "adjusted_schedule",
            "optimized_schedule",
            "new_schedule",
            "final_plan",
            "delay_adjusted_plan",
            "generated_plan",
            "progress_plan",
        )
        for key in outer_keys:
            inner = data.get(key)
            if isinstance(inner, dict):
                if "overview" in inner and "all_tasks_schedule" in inner:
                    return inner
                # 可能还有一层，递归解一次
                deeper = _unwrap_any_plan(inner)
                if deeper:
                    return deeper

        # 形式 C：再兜底扫一遍所有 dict 值，看看有没有值是 {overview, all_tasks_schedule}
        for v in data.values():
            if isinstance(v, dict) and "overview" in v and "all_tasks_schedule" in v:
                return v

        return {}

    code_match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL | re.IGNORECASE)
    if code_match:
        try:
            data = json.loads(code_match.group(1))
            unwrapped = _unwrap_any_plan(data)
            if unwrapped:
                return unwrapped
        except Exception:
            pass

    start = text.find("{")
    while start != -1:
        depth = 0
        end = -1
        in_string = False
        escape = False
        for i in range(start, len(text)):
            ch = text[i]
            if escape:
                escape = False
                continue
            if ch == "\\":
                escape = True
                continue
            if ch == '"':
                in_string = not in_string
                continue
            if in_string:
                continue
            if ch == "{":
                depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0:
                    end = i
                    break
        if end != -1:
            try:
                data = json.loads(text[start:end + 1])
                unwrapped = _unwrap_any_plan(data)
                if unwrapped:
                    return unwrapped
            except Exception:
                pass
            start = text.find("{", end + 1)
        else:
            break
    return {}


def _looks_like_plan(data: Any) -> bool:
    """兼容任意外层 key（structured_output / adjusted_plan 等）。"""
    if not isinstance(data, dict):
        return False

    def _has_core_fields(d: dict) -> bool:
        return "overview" in d and "all_tasks_schedule" in d

    if _has_core_fields(data):
        return True

    outer_keys = (
        "structured_output",
        "plan",
        "output",
        "result",
        "adjusted_plan",
        "optimized_plan",
        "updated_plan",
        "updated_schedule",
        "adjusted_schedule",
        "optimized_schedule",
        "new_schedule",
        "final_plan",
        "delay_adjusted_plan",
        "generated_plan",
        "progress_plan",
    )
    for key in outer_keys:
        inner = data.get(key)
        if isinstance(inner, dict) and _has_core_fields(inner):
            return True
    return False
